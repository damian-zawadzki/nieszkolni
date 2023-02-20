import os
import django
import re
from io import BytesIO

from django.db import connection
from django.core.files.base import ContentFile

from nieszkolni_folder.time_machine import TimeMachine
from nieszkolni_folder.cleaner import Cleaner

from nieszkolni_app.models import Material
from nieszkolni_app.models import Paper

from nieszkolni_folder.quiz_manager import QuizManager

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches
from docx.shared import Pt
from docx.shared import RGBColor

from django.http import HttpResponse

from fpdf import FPDF

os.environ["DJANGO_SETTINGS_MODULE"] = 'nieszkolni_folder.settings'
django.setup()


class DocumentManager:
    def __init__(self):
        pass

    def create_assignment_doc(
            self,
            date,
            item,
            name,
            title,
            wordcount,
            flagged_content,
            minor_errors,
            major_errors,
            reviewing_user,
            comment,
            grade
            ):

        today = TimeMachine().today()
        name_slug = name.replace(" ", "_")
        paper_title = f"A_assignment_{date}_{item}---{name_slug}.docx"

        document = Document()

        paragraph_0 = document.add_paragraph()

        paragraph_0 = document.add_heading()
        paragraph_format = paragraph_0.paragraph_format
        paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        paragraph_format.alignment

        run = paragraph_0.add_run(
            f'''
            Written by: {name}
            Submitted on: {date}
            Reviewed by: {reviewing_user}
            Item: {item}
            Number of words: {wordcount}
            Errors: {major_errors} ‰
            Punctuation and spelling errors: {minor_errors} ‰
            Grade: {grade}
            ''')
        font = run.font
        font.name = "Times New Roman"
        font.size = Pt(12)
        font.bold = False
        font.color.rgb = RGBColor(0, 0, 0)

        # Heading
        paragraph_1 = document.add_heading()
        paragraph_format = paragraph_1.paragraph_format
        paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph_format.alignment

        run = paragraph_1.add_run(title)
        font = run.font
        font.name = "Times New Roman"
        font.size = Pt(20)
        font.bold = True
        font.color.rgb = RGBColor(0, 0, 0)

        # Content
        paragraph_2 = document.add_paragraph()
        paragraph_format = paragraph_2.paragraph_format
        paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph_format.alignment

        run = paragraph_2.add_run(flagged_content)
        font = run.font
        font.name = "Times New Roman"
        font.size = Pt(12)

        # Comment
        run = paragraph_2.add_run(comment)
        font = run.font
        font.name = "Times New Roman"
        font.size = Pt(12)

        file_bytes = BytesIO()
        document.save(file_bytes)
        file_bytes.seek(0)

        paper = Paper()
        paper.stamp = TimeMachine().now_number()
        paper.title = paper_title
        paper.content.save(paper_title, ContentFile(file_bytes.read()))
        paper.save()

        return paper_title

    def create_sentences_doc(self, entries):

        list_number = str(entries[0][3])
        name = entries[0][4]
        submission_date = entries[0][5]
        reviewing_user = entries[0][6]
        item = entries[0][7]

        today = TimeMachine().today()
        name_slug = name.replace(" ", "_")
        date = TimeMachine().number_to_system_date(submission_date)
        paper_title = f"A_sentences_{date}_{item}---{name_slug}.docx"

        document = Document()

        paragraph_0 = document.add_paragraph()

        paragraph_0 = document.add_heading()
        paragraph_format = paragraph_0.paragraph_format
        paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        paragraph_format.alignment

        run = paragraph_0.add_run(
            f'''
            Translated by: {name}
            Submitted on: {date}
            Reviewed on: {today}
            Reviewed by: {reviewing_user}
            List number: {list_number}
            Item: {item}
            '''
            )
        font = run.font
        font.name = "Times New Roman"
        font.size = Pt(12)
        font.bold = False
        font.color.rgb = RGBColor(0, 0, 0)

        # Heading
        paragraph_1 = document.add_heading()
        paragraph_format = paragraph_1.paragraph_format
        paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph_format.alignment

        run = paragraph_1.add_run(f"Set: {list_number}")
        font = run.font
        font.name = "Times New Roman"
        font.size = Pt(20)
        font.bold = True
        font.color.rgb = RGBColor(0, 0, 0)

        # Content
        paragraph_2 = document.add_paragraph()

        paragraph_format = paragraph_2.paragraph_format
        paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph_format.alignment

        i = 1
        for entry in entries:
            run = paragraph_2.add_run(entry[0])
            font = run.font
            font.name = "Times New Roman"
            font.size = Pt(10)
            font.bold = False
            font.italic = True
            font.color.rgb = RGBColor(0, 0, 0)
            run.add_break()

            sentence = f"{i}. {entry[1]}"

            run = paragraph_2.add_run(sentence)
            font = run.font
            font.name = "Times New Roman"
            font.size = Pt(12)
            font.bold = True
            font.color.rgb = RGBColor(0, 0, 0)
            run.add_break()

            run = paragraph_2.add_run(entry[2])
            font = run.font
            font.name = "Times New Roman"
            font.size = Pt(12)
            font.bold = False
            font.color.rgb = RGBColor(0, 0, 0)
            run.add_break()
            run.add_break()

            i += 1

        file_bytes = BytesIO()
        document.save(file_bytes)
        file_bytes.seek(0)

        paper = Paper()
        paper.stamp = TimeMachine().now_number()
        paper.title = paper_title
        paper.content.save(paper_title, ContentFile(file_bytes.read()))
        paper.save()

        return paper_title

    def create_quizzes_doc(self, entries):

        quiz_id = entries[0][1]
        name = entries[0][3]
        submission_date = entries[0][6]
        collection_name = entries[0][8]
        result = QuizManager().display_result(quiz_id)

        name_slug = name.replace(" ", "_")
        date = TimeMachine().number_to_system_date(submission_date)
        paper_title = f"A_quiz_{date}_{quiz_id}---{name_slug}.docx"

        document = Document()

        paragraph_0 = document.add_paragraph()

        paragraph_0 = document.add_heading()
        paragraph_format = paragraph_0.paragraph_format
        paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        paragraph_format.alignment

        run = paragraph_0.add_run(
            f'''
            Filled in by: {name}
            Submitted on: {date}
            Reviewed on: {date}
            Quiz number: {quiz_id}
            Score: {result}%
            '''
            )
        font = run.font
        font.name = "Times New Roman"
        font.size = Pt(12)
        font.bold = False
        font.color.rgb = RGBColor(0, 0, 0)

        # Heading
        paragraph_1 = document.add_heading()
        paragraph_format = paragraph_1.paragraph_format
        paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph_format.alignment

        run = paragraph_1.add_run(f"Quiz: {collection_name}")
        font = run.font
        font.name = "Times New Roman"
        font.size = Pt(20)
        font.bold = True
        font.color.rgb = RGBColor(0, 0, 0)

        # Content
        paragraph_2 = document.add_paragraph()

        paragraph_format = paragraph_2.paragraph_format
        paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph_format.alignment

        i = 1
        for entry in entries:
            run = paragraph_2.add_run(entry[5])
            font = run.font
            font.name = "Times New Roman"
            font.size = Pt(10)
            font.bold = False
            font.italic = True
            font.color.rgb = RGBColor(0, 0, 0)
            run.add_break()

            sentence = f"{i}. {entry[10]}"

            run = paragraph_2.add_run(sentence)
            font = run.font
            font.name = "Times New Roman"
            font.size = Pt(12)
            font.bold = True
            font.color.rgb = RGBColor(0, 0, 0)
            run.add_break()

            run = paragraph_2.add_run(entry[4])
            font = run.font
            font.name = "Times New Roman"
            font.size = Pt(12)
            font.bold = False
            font.color.rgb = RGBColor(0, 0, 0)
            run.add_break()
            run.add_break()

            i += 1

        file_bytes = BytesIO()
        document.save(file_bytes)
        file_bytes.seek(0)

        paper = Paper()
        paper.stamp = TimeMachine().now_number()
        paper.title = paper_title
        paper.content.save(paper_title, ContentFile(file_bytes.read()))
        paper.save()

        return paper_title

    def create_timesheet_document(
            self,
            employee,
            start,
            end,
            duration,
            entries
            ):

        today = TimeMachine().today()
        employee_slug = employee.replace(" ", "_")
        paper_title = f"timesshet_{today}---{employee_slug}.docx"

        document = Document()

        paragraph_0 = document.add_paragraph()

        paragraph_0 = document.add_heading()
        paragraph_format = paragraph_0.paragraph_format
        paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph_format.alignment
        tab_stops = paragraph_0.paragraph_format.tab_stops
        tab_stop = tab_stops.add_tab_stop(Inches(2))

        run = paragraph_0.add_run(
            f'''
            Employee:\t{employee}
            Generated on:\t{today}
            From:\t{start}
            To:\t{end}
            Duration:\t{duration}
            '''
            )
        font = run.font
        font.name = "Times New Roman"
        font.size = Pt(12)
        font.bold = False
        font.color.rgb = RGBColor(0, 0, 0)

        # Heading
        paragraph_1 = document.add_heading()
        paragraph_format = paragraph_1.paragraph_format
        paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph_format.alignment

        run = paragraph_1.add_run("Timesheet")
        font = run.font
        font.name = "Times New Roman"
        font.size = Pt(20)
        font.bold = True
        font.color.rgb = RGBColor(0, 0, 0)

        # Content
        paragraph_2 = document.add_paragraph()

        paragraph_format = paragraph_2.paragraph_format
        paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph_format.alignment
        tab_stops = paragraph_2.paragraph_format.tab_stops
        tab_stop = tab_stops.add_tab_stop(Inches(1.5))
        tab_stop = tab_stops.add_tab_stop(Inches(3))
        tab_stop = tab_stops.add_tab_stop(Inches(4))

        for entry in entries:

            clock_in = entry["clock_in"]
            clock_out = entry["clock_out"]
            duration = entry["duration"]
            category_name = entry["category_name"].replace("<b>", "").replace("</b>", "")

            run = paragraph_2.add_run()
            font = run.font
            font.name = "Times New Roman"
            font.size = Pt(10)
            font.bold = False
            font.color.rgb = RGBColor(0, 0, 0)
            run.add_text(clock_in + "\t")
            run.add_text(clock_out + "\t")
            run.add_text(duration + "\t")
            run.add_text(category_name)
            run.add_break()

        file_bytes = BytesIO()
        document.save(file_bytes)
        file_bytes.seek(0)

        paper = Paper()
        paper.stamp = TimeMachine().now_number()
        paper.title = paper_title
        paper.content.save(paper_title, ContentFile(file_bytes.read()))
        paper.save()

        return paper_title

    def create_timesheet_pdf(
            self,
            employee,
            start,
            end,
            duration,
            entries
            ):

        today = TimeMachine().today()
        employee_slug = employee.replace(" ", "_")
        title = f"timesshet_{today}---{employee_slug}.pdf"

        pdf = FPDF("P", "mm", "A4")
        pdf.add_page()

        pdf.set_font("courier", "", 10)
        pdf.cell(50, 5, "Employee:", 0, 0, 1)
        pdf.cell(50, 5, f"{employee}", 0, 1, 1)
        pdf.cell(50, 5, "Generated on:", 0, 0, 1)
        pdf.cell(50, 5, f"{today}", 0, 1, 1)
        pdf.cell(50, 5, "From:", 0, 0, 1)
        pdf.cell(50, 5, f"{start}", 0, 1, 1)
        pdf.cell(50, 5, "To:", 0, 0, 1)
        pdf.cell(50, 5, f"{end}", 0, 1, 1)
        pdf.cell(50, 5, "Duration:", 0, 0, 1)
        pdf.cell(50, 5, f"{duration}", 0, 1, 1)

        pdf.set_font("courier", "B", 24)
        pdf.cell(40, 10, "", 0, 1)
        pdf.cell(40, 10, "Timesheet", 0, 1)
        pdf.cell(40, 10, "", 0, 1)

        pdf.set_font("courier", "B", 10)
        pdf.cell(45, 5, "Start", 0, 0, 1)
        pdf.cell(45, 5, "End:", 0, 0, 1)
        pdf.cell(25, 5, "Duration:", 0, 0, 1)
        pdf.cell(85, 5, "Category:", 0, 1, 1)

        pdf.set_font("courier", "", 10)
        for entry in entries:
            clock_in = entry["clock_in"]
            clock_out = entry["clock_out"]
            duration = entry["duration"]
            category_name = entry["category_name"].replace("<b>", "").replace("</b>", "")

            pdf.cell(45, 8, f"{clock_in}", 0, 0, 1)
            pdf.cell(45, 8, f"{clock_out}", 0, 0, 1)
            pdf.cell(25, 5, f"{duration}", 0, 0, 1)
            pdf.multi_cell(85, 5, f"{category_name}", 0, 1, 0)

        file = bytes(pdf.output(dest="S").encode('latin-1'))
        data = (file, title)

        return data

    def add_material(self, title, content):
        title = Cleaner().clean_quotation_marks(title)
        content = Cleaner().clean_quotation_marks(content)

        with connection.cursor() as cursor:
            cursor.execute(f'''
                INSERT INTO nieszkolni_app_material (
                title,
                content
                )
                VALUES (
                '{title}',
                '{content}'
                )
                ON CONFLICT (title)
                DO NOTHING
                ''')

    def display_materials(self):
        with connection.cursor() as cursor:
            cursor.execute(f'''
                SELECT
                title
                FROM nieszkolni_app_material
                ''')

            materials = cursor.fetchall()

            return materials

    def display_material(self, title):
        title = Cleaner().clean_quotation_marks(title)

        with connection.cursor() as cursor:
            cursor.execute(f'''
                SELECT
                title,
                content
                FROM nieszkolni_app_material
                WHERE title = '{title}'
                ''')

            material = cursor.fetchone()

            return material

    def delete_material(self, title):
        title = Cleaner().clean_quotation_marks(title[0])

        with connection.cursor() as cursor:
            cursor.execute(f'''
                DELETE FROM nieszkolni_app_material
                WHERE title = '{title}'
                ''')